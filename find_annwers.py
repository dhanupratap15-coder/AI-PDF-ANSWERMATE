
'''UPCOMING FEATURES
1. Add  more file formats (e.g., Excel, PowerPoint)
2.AI ASSISTANT FOR INTERACTIVE Q&A
3. Customizable AI Models and Parameters
4. Enhanced User Interface with Streamlit
5. SHOW DEATAILES ANSWERS
6.STUDY MODE
7.EXPORT OPTIONS'''
#import

from mistralai import Mistral
from dotenv import load_dotenv
from docx import Document
from docxcompose.composer import Composer
import os 
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
print("SmartDocs AI")

doc = Document()
load_dotenv()
client = Mistral(
    api_key=os.getenv("MISTRAL_API_KEY")
    )
print('''Give me file in form of 
      1.pdf
      2.text
      3.docx''')
input_file = input("Enter the file  : ")
input_type = input_file.split(".")[-1]
print(input_type)
if input_type == "pdf":
#load pdf
    loader = PyPDFLoader(input_file)
    documents = loader.load()

    reader = documents[0].page_content
if input_type == "text":
#load text
    loader = TextLoader(input_file)
    documents = loader.load()

    reader = documents[0].page_content
if input_type == "docx":
#load docx
    loader = Docx2txtLoader(input_file)
    documents = loader.load()

    reader = documents[0].page_content   
print('''What do you want 
      >  answer 
      >  summary
      >  ASK Questions
      ''')
input_ask = input("Enter your choice : ")
if input_ask.lower() == "answer":
    #split question 
    lines = reader.split("\n")
    questions = []
    for line in lines :
        line = line.strip()
        for i,h in enumerate(lines):
            if line.startswith(str(i) + "."):
                q = line.split(".",1)[1].strip()
    
                questions.append(q)

    # Use AI
    # import requests
    # url = "http://localhost:11434/api/generate"
    # print("Making....")
    file_name = []
    doc.add_heading("AI Generated Answers", level=1)
    print("you want marge all docx files ?")   
    marge = input("yes or no : ")  
    for i,pro in enumerate(questions):
        prompt = f"""
        Question: {pro}
        Write a clear, well-structured answer with  key points,
        examples, and conclusion  using
        English style for easy understanding.
        """
        print(f"{i+1}. {pro}")
        # response = requests.post(
        #     url,
        #     json={
        #         "model": "llama3.2",
        #         "prompt": prompt,
        #         "stream": False
        #     }
        # )
        # data = response.json()['response']

        # print(data)    
        ####USE MISTRAL API 
    
        response = client.chat.complete(
            model="mistral-small-latest" ,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        opt= response.choices[0].message.content
        doc.add_paragraph(opt)
        if marge.lower() == "no":
            doc.save(f"output{i+1}.docx") 
        file_name.append(f"output{i+1}.docx") 
        print(f"{i+1} . successfully")  
      
    if marge.lower() == "yes":
    # Open the first document
        main_doc = Document(file_name[0])    
        composer = Composer(main_doc)
            # Append each document
        for file in file_name[1:]:
            doc = Document(file)
            composer.append(doc)
        # Save merged file
        composer.save("merged_output.docx")

        print("DOCX files merged successfully!")
if input_ask.lower() == "summary":
    prompt = f"""
    Summarize the following content in a clear and concise manner, highlighting the key points and main ideas:
    {reader}
    """
    response = client.chat.complete(
        model="mistral-small-latest" ,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )
    opt= response.choices[0].message.content
    print(opt)       
if input_ask.lower() == "ask questions":
    ask = input("Enter your question : ")
    prompt = f"""
    Based on the following content,{ask}
    {reader}
    """
    response = client.chat.complete(
        model="mistral-small-latest" ,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )
    opt= response.choices[0].message.content
    print(opt)     