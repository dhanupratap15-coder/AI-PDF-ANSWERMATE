#import
from mistralai import Mistral
from dotenv import load_dotenv
from docx import Document
from docxcompose.composer import Composer
import os 
from docx import Document
doc = Document()
load_dotenv()
client = Mistral(
    api_key=os.getenv("MISTRAL_API_KEY")
    )
from PyPDF2 import PdfReader
#load pdf
reader = PdfReader("sample1.pdf")
# show text
for i , page in enumerate(reader.pages):
    text=page.extract_text()
    print(f"--page{i+1}--")
#split question 
lines = text.split("\n")
questions = []
for line in lines :
    line = line.strip()
    for i,h in enumerate(lines):
        if line.startswith(str(i) + "."):
            q = line.split(".",1)[1].strip()
   
            questions.append(q)

# Use AI
# import requests
# url = "http://localhost:11434/api/generate"
# print("Making....")
file_name = []
doc.add_heading("AI Generated Answers", level=1)
for i,pro in enumerate(questions):
    prompt = f"""
    Question: {pro}
    Write a clear, well-structured answer with 10 key points,
    examples, and conclusion  using
    English style for easy understanding.
    """
    print(f"{i+1}. {pro}")
    # response = requests.post(
    #     url,
    #     json={
    #         "model": "llama3.2",
    #         "prompt": prompt,
    #         "stream": False
    #     }
    # )
    # data = response.json()['response']

    # print(data)    
    ####USE MISTRAL API 
   
    response = client.chat.complete(
        model="mistral-small-latest" ,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )
    opt= response.choices[0].message.content
    doc.add_paragraph(opt)
    
    doc.save(f"output{i+1}.docx") 
    file_name.append(f"output{i+1}.docx") 
    print(f"{i+1} . successfully")  
print("you want marge all docx files ?")   
marge = input("yes or no : ")  
if marge.lower() == "yes":
# Open the first document
    main_doc = Document(file_name[0])    
    composer = Composer(main_doc)
        # Append each document
    for file in file_name[1:]:
        doc = Document(file)
        composer.append(doc)
    # Save merged file
    composer.save("merged_output.docx")

    print("DOCX files merged successfully!")
        
        
